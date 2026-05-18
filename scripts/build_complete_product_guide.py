from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "artifacts"
ASSET_DIR = OUT_DIR / "guide_assets"
DOCX_PATH = OUT_DIR / "Autonomous_Security_Lab_Assistant_Complete_Guide.docx"

BLUE = RGBColor(31, 78, 121)
TEAL = RGBColor(0, 112, 112)
DARK = RGBColor(36, 45, 57)
GRAY = RGBColor(92, 102, 112)
LIGHT_BLUE = "DDEBF7"
LIGHT_TEAL = "DDEFEF"
LIGHT_GRAY = "F3F6F8"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = create_diagrams()

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    add_section(
        doc,
        "1. Executive Summary",
        [
            "The Autonomous Security Lab Assistant is a safe-by-default cybersecurity lab assistant built around MCP-style tool orchestration. It lets a user or AI agent inspect authorized lab targets, collect structured evidence, generate findings, save reports, export SARIF, index runs, and maintain tamper-evident audit logs.",
            "The project is intentionally not a chatbot, not a CRUD app, and not a simple API wrapper. It is a security orchestration backend that demonstrates AI agents, MCP architecture, defensive cybersecurity workflows, policy enforcement, persistence, reporting, and backend engineering discipline.",
            "The central design rule is simple: policy first, tools second, evidence always.",
        ],
    )

    add_callout(
        doc,
        "Plain-English Definition",
        "This project is a secure AI-agent tool server for authorized cybersecurity lab reconnaissance. It checks scope before action, refuses unsafe requests, runs bounded tools, and records evidence.",
    )

    add_section(
        doc,
        "2. Who Uses This Product",
        [
            "Cybersecurity students use it to learn safe reconnaissance workflows without building unsafe habits.",
            "Backend engineers use it to study protocol design, tool registries, persistence, reporting, and secure input validation.",
            "AI engineers use it to understand how an agent can call tools through a controlled MCP-style interface.",
            "Security researchers and CTF players use it inside owned labs, private ranges, and intentionally vulnerable training environments.",
            "Teachers and mentors use it as a controlled demonstration of how security automation should be bounded.",
        ],
    )

    add_table(
        doc,
        "User groups and value",
        ["User", "Why they use it", "What they learn or gain"],
        [
            ["Cybersecurity student", "Practice safe lab recon", "Scope validation, evidence, reports"],
            ["Backend engineer", "Study serious service design", "JSON-RPC, validation, storage, tests"],
            ["AI engineer", "Build safe agent tools", "MCP-style tool schemas and orchestration"],
            ["Security team", "Prototype defensive workflows", "Auditability, SARIF, risk scoring"],
            ["Portfolio builder", "Show advanced capability", "Security plus backend plus agent architecture"],
        ],
    )

    add_picture(doc, diagrams["architecture"], "Figure 1. Product architecture: clients call a controlled tool server, tools pass through policy, and workflow artifacts are persisted.")

    add_section(
        doc,
        "3. Primary Use Case",
        [
            "The main use case is authorized lab reconnaissance. A user points the assistant at an allowed local or private target, chooses ports, and receives structured results.",
            "For example, a developer may run a local test application on 127.0.0.1:8000 and ask the assistant to inspect common web ports. The assistant validates scope, scans only the requested ports, collects HTTP headers, checks TLS metadata when HTTPS exists, checks security.txt and robots.txt, extracts safe HTML page intelligence, generates findings, scores risk, and saves artifacts.",
        ],
    )

    add_code(
        doc,
        "CLI recon example",
        "python -m security_lab_assistant.cli recon 127.0.0.1 --ports 80,443,8000,8080",
    )

    add_picture(doc, diagrams["workflow"], "Figure 2. Autonomous recon workflow from input validation to persisted artifacts.")

    add_section(
        doc,
        "4. What The Product Does",
        [
            "The product exposes narrow security tools. Each tool has a name, description, input schema, and handler function. The MCP-style server lists these tools and allows clients to call them with JSON arguments.",
            "Every active network tool checks the policy engine before it touches a target. Unsafe requests produce structured refusals instead of silent failures or uncontrolled exceptions.",
        ],
    )

    add_table(
        doc,
        "Tool catalog",
        ["Tool", "Purpose", "Safety controls"],
        [
            ["scope.validate", "Checks whether a target is allowed", "Target normalization, CIDR and hostname policy"],
            ["scan.tcp_connect", "Bounded TCP connect scan", "Port limit, blocked ports, timeout, worker cap"],
            ["recon.http_headers", "Collect selected HTTP headers", "URL policy, no auto-follow redirects"],
            ["recon.tls_certificate", "Inspect TLS certificate metadata", "Target and port policy"],
            ["recon.well_known_security", "Check robots.txt and security.txt", "In-scope URL checks"],
            ["recon.web_page_intel", "Extract safe HTML indicators", "No JavaScript execution"],
            ["web.fetch_text", "Fetch bounded text content", "Max byte limit, URL policy"],
            ["analyze.security_headers", "Create header-hardening finding", "Structured analysis only"],
            ["run.list", "List persisted runs", "Bounded limit"],
            ["run.get", "Load a run by ID", "UUID-only lookup"],
            ["run.search", "Search indexed runs", "Bounded query, status validation"],
            ["run.verify_audit", "Verify audit hash chain", "Tamper-evidence check"],
            ["workflow.autonomous_recon", "Run the full workflow", "All controls combined"],
        ],
    )

    add_section(
        doc,
        "5. How To Run The Product",
        [
            "Open PowerShell in the project directory. In this workspace, the project lives at D:\\MCP_Demo.",
            "The project requires Python 3.11 or newer and currently has no third-party runtime dependencies beyond the bundled environment used to generate this guide.",
        ],
    )

    add_code(doc, "Run all tests", "python -m unittest discover -s tests")
    add_code(doc, "Run a recon workflow", "python -m security_lab_assistant.cli recon 127.0.0.1 --ports 80,443,8000,8080")
    add_code(doc, "List saved runs", "python -m security_lab_assistant.cli runs")
    add_code(doc, "Search saved runs", "python -m security_lab_assistant.cli search --query 127.0.0.1 --status completed")
    add_code(doc, "Show one saved run", "python -m security_lab_assistant.cli show <run_id>")

    add_section(
        doc,
        "6. How To Use The MCP-Style Server",
        [
            "The server reads newline-delimited JSON-RPC messages from stdin and writes JSON-RPC responses to stdout. This keeps the transport simple and inspectable.",
            "Supported methods are initialize, tools/list, tools/call, and notifications/initialized.",
        ],
    )
    add_code(doc, "Start the MCP-style server", "python -m security_lab_assistant.server")
    add_code(doc, "Initialize request", '{"jsonrpc":"2.0","id":1,"method":"initialize","params":{"clientInfo":{"name":"demo-client","version":"0.1.0"}}}')
    add_code(doc, "List tools request", '{"jsonrpc":"2.0","id":2,"method":"tools/list","params":{}}')
    add_code(doc, "Run workflow request", '{"jsonrpc":"2.0","id":4,"method":"tools/call","params":{"name":"workflow.autonomous_recon","arguments":{"target":"127.0.0.1","ports":[80,443,8000,8080],"objective":"baseline product smoke test"}}}')

    add_section(
        doc,
        "7. Security Model",
        [
            "The security model is strict because this project is used in a cybersecurity context. The product is built to refuse risky behavior by default.",
            "No software can honestly promise zero vulnerabilities. The correct professional claim is that the project uses layered controls, regression tests, narrow tool boundaries, and explicit residual-risk documentation.",
        ],
    )
    add_picture(doc, diagrams["security"], "Figure 3. Security boundary model: every active operation must pass policy before reaching a target or artifact path.")
    add_table(
        doc,
        "Security controls",
        ["Control", "What it prevents"],
        [
            ["Deny-by-default DNS targets", "Accidental public-host scans"],
            ["CIDR allowlist", "Out-of-scope IP targets"],
            ["Hostname allowlist", "Unapproved DNS names"],
            ["Blocked ports", "Unsafe or unwanted service probing"],
            ["Max ports per scan", "Large unintended scans"],
            ["URL credential rejection", "Credential leakage and ambiguous authority"],
            ["Fragment and parameter rejection", "Ambiguous browser-only or legacy URL behavior"],
            ["UUID-only run lookup", "Path traversal through run IDs"],
            ["Project-local artifact root", "Writing outside the workspace"],
            ["Atomic writes", "Partial or corrupted artifact files"],
            ["Hash-chained audit log", "Undetected audit tampering"],
            ["No shell tool", "Arbitrary command execution"],
            ["No JavaScript execution", "Browser-side execution risk"],
        ],
    )

    add_section(
        doc,
        "8. Artifact System",
        [
            "A serious security product must preserve evidence. This project creates a local artifact directory when workflows run.",
            "The artifact directory is project-local by policy. Absolute artifact directories and out-of-project paths are refused.",
        ],
    )
    add_picture(doc, diagrams["artifacts"], "Figure 4. Artifact layout created by workflow runs.")
    add_table(
        doc,
        "Artifacts",
        ["Path", "Purpose"],
        [
            [".security_lab_assistant/runs/<run_id>.json", "Complete structured evidence"],
            [".security_lab_assistant/reports/<run_id>.md", "Human-readable Markdown report"],
            [".security_lab_assistant/exports/<run_id>.sarif.json", "Security tooling export"],
            [".security_lab_assistant/audit/events.jsonl", "Tamper-evident audit events"],
            [".security_lab_assistant/index.sqlite3", "Run and audit index"],
        ],
    )

    add_section(
        doc,
        "9. File-By-File Walkthrough",
        [
            "This section explains what each important file does and why it exists.",
        ],
    )

    add_file_walkthrough(doc)

    add_section(
        doc,
        "10. Line-By-Line Conceptual Walkthrough Of The Main Workflow",
        [
            "The main workflow lives in security_lab_assistant/workflows/autonomous_recon.py. The exact source code should be read in the repository, but the following walkthrough explains the logic line by line at a conceptual level.",
        ],
    )
    add_code_walkthrough(doc)

    add_section(
        doc,
        "11. How Testing Works",
        [
            "The tests are built with Python's standard unittest framework. The suite verifies policy behavior, refusal paths, protocol handling, persistence, SARIF export, risk scoring, audit verification, and security edge cases.",
        ],
    )
    add_table(
        doc,
        "Test files",
        ["Test file", "What it verifies"],
        [
            ["tests/test_policy.py", "Default policy, local targets, public IP refusal, port limits"],
            ["tests/test_tools.py", "Tool refusal behavior and workflow stopping"],
            ["tests/test_mcp_protocol.py", "JSON-RPC initialize, tools/list, tools/call validation"],
            ["tests/test_product_edges.py", "Persistence, invalid URLs, DNS denial, path traversal, limits"],
            ["tests/test_enterprise_features.py", "SARIF, risk scoring, SQLite search, audit verification"],
        ],
    )
    add_code(doc, "Expected test command", "python -m unittest discover -s tests")
    add_code(doc, "Expected result", "Ran 28 tests\nOK")

    add_section(
        doc,
        "12. Example Output Explained",
        [
            "A successful workflow response contains ok, run_id, target, status, open ports, severity counts, risk information, warnings, evidence steps, findings, and artifact paths.",
        ],
    )
    add_code(
        doc,
        "Simplified output",
        '{\n  "ok": true,\n  "data": {\n    "run_id": "44da8f18-1368-425a-85d8-ee1e519e6a2f",\n    "target": "127.0.0.1",\n    "status": "completed",\n    "open_ports": [],\n    "risk": {"score": 2, "band": "low"},\n    "warnings": ["No open ports were detected in the requested scan set."],\n    "run_path": ".security_lab_assistant/runs/<run_id>.json",\n    "report_path": ".security_lab_assistant/reports/<run_id>.md",\n    "sarif_path": ".security_lab_assistant/exports/<run_id>.sarif.json"\n  }\n}',
    )

    add_section(
        doc,
        "13. What This Product Does Not Do",
        [
            "It does not exploit vulnerabilities.",
            "It does not brute force passwords.",
            "It does not run malware.",
            "It does not execute arbitrary shell commands.",
            "It does not scan arbitrary public internet targets by default.",
            "It does not bypass authentication.",
            "It does not perform stealth scanning.",
            "It does not replace professional security testing.",
        ],
    )

    add_section(
        doc,
        "14. How To Present This Project",
        [
            "Long version: Autonomous Security Lab Assistant is a safe-by-default MCP-style security orchestration backend. It validates scope before action, exposes narrow security tools through structured schemas, performs bounded recon against authorized lab targets, persists evidence, generates reports, exports SARIF, maintains a SQLite run index, and records tamper-evident audit logs.",
            "Short version: It is a secure AI-agent tool server for authorized cybersecurity lab reconnaissance.",
        ],
    )

    add_section(
        doc,
        "15. Future Improvements",
        [
            "Add official MCP SDK integration.",
            "Add authenticated HTTP transport if the server is exposed beyond stdio.",
            "Add Docker Compose vulnerable lab targets for demos.",
            "Add HTML report export.",
            "Add CI with unit tests, type checks, and security scans.",
            "Add signed releases and stricter filesystem permission checks.",
            "Add role-based policies and human approval gates for higher-risk tools.",
        ],
    )

    add_callout(
        doc,
        "Final Summary",
        "This project combines AI agents, MCP architecture, cybersecurity guardrails, backend engineering, security reporting, persistent evidence, auditability, and tests. The most important principle is: policy first, tools second, evidence always.",
    )

    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 28, BLUE),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autonomous Security Lab Assistant")
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Complete Product Guide")
    run.font.name = "Aptos"
    run.font.size = Pt(18)
    run.font.color.rgb = TEAL

    add_spacer(doc, 0.25)
    add_callout(
        doc,
        "Product Purpose",
        "A secure MCP-style backend for authorized cybersecurity lab reconnaissance, evidence collection, reporting, auditability, and AI-agent tool orchestration.",
    )
    add_table(
        doc,
        "Document metadata",
        ["Field", "Value"],
        [
            ["Project", "Autonomous Security Lab Assistant"],
            ["Version", "0.4.0"],
            ["Workspace", "D:\\MCP_Demo"],
            ["Audience", "Students, backend engineers, AI engineers, security learners, reviewers"],
            ["Generated artifact", "PDF manual with diagrams and detailed explanations"],
        ],
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    items = [
        "Executive Summary",
        "Who Uses This Product",
        "Primary Use Case",
        "What The Product Does",
        "How To Run The Product",
        "How To Use The MCP-Style Server",
        "Security Model",
        "Artifact System",
        "File-By-File Walkthrough",
        "Line-By-Line Conceptual Workflow",
        "How Testing Works",
        "Example Output Explained",
        "What This Product Does Not Do",
        "How To Present This Project",
        "Future Improvements",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_page_break()


def add_section(doc: Document, title: str, paragraphs: list[str]) -> None:
    doc.add_heading(title, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_TEAL)
    set_cell_margins(cell, 160, 160, 160, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = TEAL
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.bold = True
    r.font.color.rgb = DARK

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell, 120, 120, 100, 100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell, 120, 120, 100, 100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].add_run(value)
    doc.add_paragraph()


def add_code(doc: Document, title: str, code: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_margins(cell, 120, 120, 120, 120)
    for idx, line in enumerate(code.splitlines() or [""]):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.65))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


def add_file_walkthrough(doc: Document) -> None:
    rows = [
        ["pyproject.toml", "Package metadata, version, Python requirement, CLI entry points."],
        ["configs/lab_policy.json", "Default security policy: allowed CIDRs, hostnames, blocked ports, scan limits, artifact path."],
        ["models.py", "Core dataclasses: ToolResult, Finding, RunContext."],
        ["policy.py", "Security gatekeeper for targets, URLs, ports, and artifact paths."],
        ["validation.py", "Shared safe parsing for strings, ports, limits, statuses, run IDs, and timeouts."],
        ["tools/base.py", "ToolSpec abstraction and structured refusal helper."],
        ["tools/scope.py", "Implements scope.validate."],
        ["tools/network.py", "Implements TCP scan, HTTP headers, TLS certs, well-known files, fetch, and page intel."],
        ["tools/reporting.py", "Analyzes security headers and creates findings."],
        ["tools/runs.py", "Run list, get, search, and audit verification tools."],
        ["tools/registry.py", "Central catalog of exposed MCP tools and schemas."],
        ["workflows/autonomous_recon.py", "Main orchestration loop from validation to persistence."],
        ["storage.py", "Atomic artifact writes, SQLite index, audit hash chain, run persistence."],
        ["reporting.py", "Markdown report renderer."],
        ["risk.py", "Risk scoring and band calculation."],
        ["sarif.py", "SARIF export renderer."],
        ["mcp_protocol.py", "MCP-style JSON-RPC method handler."],
        ["server.py", "Stdio JSON-RPC server."],
        ["cli.py", "Command-line interface for recon and run management."],
        ["tests/", "Security, protocol, persistence, and feature regression tests."],
    ]
    add_table(doc, "Important files", ["File", "Purpose"], rows)


def add_code_walkthrough(doc: Document) -> None:
    rows = [
        ["1", "Read target", "require_string rejects empty, overlong, or control-character input."],
        ["2", "Read objective", "Uses provided objective or defaults to baseline web reconnaissance."],
        ["3", "Parse ports", "parse_ports rejects invalid values, booleans, duplicates, and unsafe types."],
        ["4", "Create RunContext", "A run object stores target, status, phases, evidence, findings, warnings."],
        ["5", "Audit start", "workflow.started is appended to the tamper-evident audit log."],
        ["6", "Scope phase", "scope.validate confirms target is allowed before network activity."],
        ["7", "Refusal path", "If scope fails, status becomes refused and artifacts are still saved."],
        ["8", "Port scan phase", "scan.tcp_connect checks requested ports with bounded concurrency."],
        ["9", "Service recon phase", "Open web ports are converted into HTTP or HTTPS URLs."],
        ["10", "Header recon", "recon.http_headers captures selected headers without following redirects."],
        ["11", "TLS recon", "HTTPS services trigger certificate metadata collection."],
        ["12", "Well-known checks", "robots.txt and security.txt locations are inspected."],
        ["13", "Page intelligence", "HTML structure is parsed safely without executing scripts."],
        ["14", "Analysis", "analyze.security_headers turns observations into findings."],
        ["15", "Risk and reports", "Risk score, Markdown report, SARIF, JSON evidence, SQLite index, and audit events are written."],
        ["16", "Return result", "The workflow returns run ID, risk, warnings, findings, steps, and artifact paths."],
    ]
    add_table(doc, "Conceptual line-by-line workflow", ["Step", "Code idea", "Explanation"], rows)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Autonomous Security Lab Assistant Complete Product Guide")
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY


def add_spacer(doc: Document, inches: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(inches * 72)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int, bottom: int, left: int, right: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def create_diagrams() -> dict[str, Path]:
    return {
        "architecture": architecture_diagram(),
        "workflow": workflow_diagram(),
        "security": security_diagram(),
        "artifacts": artifacts_diagram(),
    }


def font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str, outline: str = "#2f3b4a") -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = wrap_text(text, 24)
    line_height = 24
    total = line_height * len(lines)
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font(19, True))
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill="#17212b", font=font(19, True))
        y += line_height


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#365f91", width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        points = [(x2, y2), (x2 - 12, y2 - 8), (x2 - 12, y2 + 8)]
    else:
        points = [(x2, y2), (x2 + 12, y2 - 8), (x2 + 12, y2 + 8)]
    draw.polygon(points, fill="#365f91")


def diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1400, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1400, 92), fill="#1f4e79")
    draw.text((42, 26), title, fill="white", font=font(34, True))
    return img, draw


def architecture_diagram() -> Path:
    img, draw = diagram_canvas("Product Architecture")
    boxes = [
        ((70, 170, 300, 270), "CLI or MCP Client", "#ddebf7"),
        ((390, 170, 620, 270), "JSON-RPC Server", "#ddebf7"),
        ((710, 170, 940, 270), "Tool Registry", "#ddebf7"),
        ((1030, 170, 1260, 270), "Policy Engine", "#fce4d6"),
        ((390, 390, 620, 490), "Autonomous Workflow", "#ddefff"),
        ((710, 390, 940, 490), "Recon Tools", "#e2f0d9"),
        ((1030, 390, 1260, 490), "Artifacts + Audit", "#e4dfec"),
    ]
    for box in boxes:
        draw_box(draw, *box)
    for start, end in [((300, 220), (390, 220)), ((620, 220), (710, 220)), ((940, 220), (1030, 220)), ((505, 270), (505, 390)), ((620, 440), (710, 440)), ((940, 440), (1030, 440)), ((825, 390), (825, 270))]:
        arrow(draw, start, end)
    path = ASSET_DIR / "architecture.png"
    img.save(path)
    return path


def workflow_diagram() -> Path:
    img, draw = diagram_canvas("Autonomous Recon Workflow")
    labels = ["Input Validation", "Scope Check", "TCP Scan", "Web Recon", "Analysis", "Risk Score", "Reports", "Audit + Index"]
    x = 50
    y = 190
    w = 275
    h = 92
    for idx, label in enumerate(labels):
        row = idx // 4
        col = idx % 4
        x1 = 65 + col * 330
        y1 = 170 + row * 230
        draw_box(draw, (x1, y1, x1 + w, y1 + h), label, "#e2f0d9" if idx not in {1, 7} else "#fce4d6")
        if col < 3:
            arrow(draw, (x1 + w, y1 + h // 2), (x1 + 330, y1 + h // 2))
        elif row == 0:
            arrow(draw, (x1 + w // 2, y1 + h), (x1 + w // 2, y1 + 230))
    path = ASSET_DIR / "workflow.png"
    img.save(path)
    return path


def security_diagram() -> Path:
    img, draw = diagram_canvas("Security Boundary Model")
    draw_box(draw, (80, 190, 310, 300), "Untrusted Request", "#f8cbad")
    draw_box(draw, (430, 190, 700, 300), "Validation Layer", "#ddebf7")
    draw_box(draw, (820, 190, 1090, 300), "Policy Engine", "#fce4d6")
    draw_box(draw, (430, 430, 700, 540), "Allowed Tool Action", "#e2f0d9")
    draw_box(draw, (820, 430, 1090, 540), "Structured Refusal", "#f4cccc")
    arrow(draw, (310, 245), (430, 245))
    arrow(draw, (700, 245), (820, 245))
    arrow(draw, (955, 300), (565, 430))
    arrow(draw, (955, 300), (955, 430))
    draw.text((1120, 210), "Everything must pass policy.\nUnsafe input becomes a refusal.\nNo shell execution tool exists.\nAudit events record workflow actions.", fill="#263238", font=font(24))
    path = ASSET_DIR / "security.png"
    img.save(path)
    return path


def artifacts_diagram() -> Path:
    img, draw = diagram_canvas("Artifact System")
    draw_box(draw, (80, 180, 360, 290), ".security_lab_assistant", "#ddebf7")
    children = [
        ((500, 140, 800, 220), "runs/*.json\nStructured Evidence", "#e2f0d9"),
        ((500, 250, 800, 330), "reports/*.md\nHuman Report", "#e2f0d9"),
        ((500, 360, 800, 440), "exports/*.sarif.json\nSecurity Export", "#e2f0d9"),
        ((500, 470, 800, 550), "audit/events.jsonl\nHash Chain", "#fce4d6"),
        ((930, 305, 1230, 405), "index.sqlite3\nSearch + Metadata", "#e4dfec"),
    ]
    for xy, text, fill in children:
        draw_box(draw, xy, text, fill)
        arrow(draw, (360, 235), (xy[0], (xy[1] + xy[3]) // 2))
    arrow(draw, (800, 400), (930, 355))
    path = ASSET_DIR / "artifacts.png"
    img.save(path)
    return path


def wrap_text(text: str, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        if sum(len(part) for part in current) + len(current) + len(word) > width and current:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


if __name__ == "__main__":
    main()
